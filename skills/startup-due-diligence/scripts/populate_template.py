#!/usr/bin/env python3
"""
Deterministic DOCX Template Population Script

This script performs mechanical placeholder replacement in Word documents.
It ONLY touches {{PLACEHOLDER}} patterns - no AI or heuristic modifications.

Usage:
    python populate_template.py <template_path> <data_json_path> <output_path>

Example:
    python populate_template.py templates/doc-request.docx data.json output.docx
"""

import json
import sys
import re
import copy
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx --break-system-packages")
    sys.exit(1)


def find_placeholders(text: str) -> list[str]:
    """Find all {{PLACEHOLDER}} patterns in text."""
    return re.findall(r'\{\{([A-Z][A-Z0-9_]*)\}\}', text)


def replace_in_paragraph(paragraph, data: dict) -> int:
    """
    Replace placeholders in a paragraph while preserving formatting.
    Returns count of replacements made.
    """
    replacements = 0

    # First, check if placeholder spans multiple runs (common in Word)
    full_text = paragraph.text
    placeholders = find_placeholders(full_text)

    if not placeholders:
        return 0

    # Simple case: placeholder is contained in a single run
    for run in paragraph.runs:
        for placeholder in placeholders:
            pattern = f"{{{{{placeholder}}}}}"
            if pattern in run.text:
                value = data.get(placeholder, pattern)  # Keep placeholder if no data
                run.text = run.text.replace(pattern, str(value))
                replacements += 1

    # Complex case: placeholder spans multiple runs
    # Rebuild paragraph text if placeholders still exist
    if find_placeholders(paragraph.text):
        full_text = paragraph.text
        for placeholder in placeholders:
            pattern = f"{{{{{placeholder}}}}}"
            value = data.get(placeholder, pattern)
            full_text = full_text.replace(pattern, str(value))

        # Clear all runs and set text on first run (loses some formatting)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            for run in paragraph.runs[1:]:
                run.text = ""
            first_run.text = full_text
            replacements += len(placeholders)

    return replacements


def replace_in_table(table, data: dict) -> int:
    """Replace placeholders in all cells of a table."""
    replacements = 0
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replacements += replace_in_paragraph(paragraph, data)
    return replacements


def replace_in_headers_footers(doc, data: dict) -> int:
    """Replace placeholders in headers and footers."""
    replacements = 0
    for section in doc.sections:
        # Headers
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragraph in header.paragraphs:
                    replacements += replace_in_paragraph(paragraph, data)
                for table in header.tables:
                    replacements += replace_in_table(table, data)

        # Footers
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragraph in footer.paragraphs:
                    replacements += replace_in_paragraph(paragraph, data)
                for table in footer.tables:
                    replacements += replace_in_table(table, data)

    return replacements


def populate_template(template_path: str, data: dict, output_path: str) -> dict:
    """
    Populate a DOCX template with data.

    Args:
        template_path: Path to the template .docx file
        data: Dictionary mapping placeholder names to values
        output_path: Path for the output .docx file

    Returns:
        Dictionary with stats about the operation
    """
    # Validate inputs
    template_path = Path(template_path)
    output_path = Path(output_path)

    if not template_path.exists():
        raise FileNotFoundError(f"Template not found: {template_path}")

    if not template_path.suffix.lower() == '.docx':
        raise ValueError(f"Template must be a .docx file: {template_path}")

    # Load document
    doc = Document(template_path)

    # Track statistics
    stats = {
        "template": str(template_path),
        "output": str(output_path),
        "replacements": 0,
        "placeholders_found": set(),
        "placeholders_filled": set(),
        "placeholders_unfilled": set()
    }

    # Find all placeholders first
    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)

    full_text = " ".join(all_text)
    stats["placeholders_found"] = set(find_placeholders(full_text))

    # Perform replacements
    for paragraph in doc.paragraphs:
        stats["replacements"] += replace_in_paragraph(paragraph, data)

    for table in doc.tables:
        stats["replacements"] += replace_in_table(table, data)

    stats["replacements"] += replace_in_headers_footers(doc, data)

    # Calculate filled vs unfilled
    stats["placeholders_filled"] = stats["placeholders_found"] & set(data.keys())
    stats["placeholders_unfilled"] = stats["placeholders_found"] - set(data.keys())

    # Convert sets to lists for JSON serialization
    stats["placeholders_found"] = sorted(list(stats["placeholders_found"]))
    stats["placeholders_filled"] = sorted(list(stats["placeholders_filled"]))
    stats["placeholders_unfilled"] = sorted(list(stats["placeholders_unfilled"]))

    # Ensure output directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Save document
    doc.save(output_path)

    return stats


def main():
    """CLI entry point."""
    if len(sys.argv) != 4:
        print("Usage: python populate_template.py <template.docx> <data.json> <output.docx>")
        print("\nExample:")
        print("  python populate_template.py templates/doc-request.docx data.json output.docx")
        sys.exit(1)

    template_path = sys.argv[1]
    data_path = sys.argv[2]
    output_path = sys.argv[3]

    # Load data
    try:
        with open(data_path, 'r') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"ERROR: Data file not found: {data_path}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON in data file: {e}")
        sys.exit(1)

    # Populate template
    try:
        stats = populate_template(template_path, data, output_path)

        print(f"SUCCESS: Template populated")
        print(f"  Output: {stats['output']}")
        print(f"  Replacements: {stats['replacements']}")
        print(f"  Placeholders filled: {len(stats['placeholders_filled'])}")

        if stats['placeholders_unfilled']:
            print(f"  WARNING: Unfilled placeholders: {stats['placeholders_unfilled']}")

    except Exception as e:
        print(f"ERROR: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
