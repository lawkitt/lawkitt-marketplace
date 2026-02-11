#!/usr/bin/env python3
"""
Create placeholder DOCX templates for the due diligence skill.
These templates use {{PLACEHOLDER}} syntax for deterministic population.

Run this script to generate initial templates that can be replaced with
professionally designed versions.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


def add_heading(doc, text, level=1):
    """Add a heading to the document."""
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False):
    """Add a paragraph to the document."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def create_document_request_list(output_path: Path):
    """Create the document request list template."""
    doc = Document()

    # Title
    title = doc.add_heading('Due Diligence Document Request List', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header info
    doc.add_paragraph()
    add_paragraph(doc, 'To: {{COMPANY_NAME}}')
    add_paragraph(doc, 'From: {{INVESTOR_NAME}}')
    add_paragraph(doc, 'Date: {{REQUEST_DATE}}')
    add_paragraph(doc, 'Re: {{INVESTMENT_TYPE}} Due Diligence')

    doc.add_paragraph()
    add_paragraph(doc, 'In connection with the proposed investment, please provide the following documents and information. Please respond by {{RESPONSE_DEADLINE}}.')

    # Section 1: Corporate
    add_heading(doc, '1. Corporate Formation & Governance', 2)
    items = [
        'Certificate of Incorporation (and all amendments)',
        'Bylaws (current)',
        'Good Standing Certificate (state of incorporation)',
        'Foreign qualification certificates',
        'All Board meeting minutes and written consents',
        'All Stockholder meeting minutes and written consents',
        'Organizational chart'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Section 2: Cap Table
    add_heading(doc, '2. Capitalization & Securities', 2)
    items = [
        'Current cap table (fully diluted)',
        'Stock ledger',
        'All SAFE agreements',
        'All convertible note agreements',
        'Prior financing documents (if any)',
        'Equity incentive plan and all amendments',
        'Stock option agreements (all grants)',
        '409A valuation reports',
        'Form D filings'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Section 3: Founder/Employee
    add_heading(doc, '3. Founders & Employees', 2)
    items = [
        'Founder stock purchase agreements',
        '83(b) elections',
        'Employment agreements for founders and key employees',
        'PIIA / CIIAA agreements (all employees)',
        'Contractor agreements',
        'Employee handbook'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Section 4: IP
    add_heading(doc, '4. Intellectual Property', 2)
    items = [
        'IP assignment agreements (founders)',
        'IP assignment agreements (contractors)',
        'Patent filings and registrations',
        'Trademark registrations',
        'Domain name registrations',
        'Open source software audit/inventory',
        'Third-party license agreements'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Section 5: Contracts
    add_heading(doc, '5. Material Contracts', 2)
    items = [
        'Top 10 customer contracts',
        'Key vendor/supplier agreements',
        'Partnership agreements',
        'Loan agreements and debt instruments',
        'Lease agreements',
        'Insurance policies'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Section 6: Compliance
    add_heading(doc, '6. Compliance & Legal', 2)
    items = [
        'Privacy policy',
        'Terms of service',
        'Any regulatory licenses or permits',
        'Pending or threatened litigation',
        'Settlement agreements'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Section 7: Financial
    add_heading(doc, '7. Financial', 2)
    items = [
        'Financial statements (historical)',
        'Financial projections',
        'Tax returns (3 years)',
        'Bank statements (6 months)'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Footer
    doc.add_paragraph()
    add_paragraph(doc, '{{ADDITIONAL_NOTES}}')

    doc.add_paragraph()
    add_paragraph(doc, 'Contact: {{INVESTOR_CONTACT}} ({{INVESTOR_EMAIL}})')

    doc.save(output_path)
    print(f"Created: {output_path}")


def create_diligence_report(output_path: Path):
    """Create the diligence report template."""
    doc = Document()

    # Title
    title = doc.add_heading('Legal Due Diligence Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('{{COMPANY_NAME}}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Meta info
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    cells = [
        ('Report Date:', '{{REPORT_DATE}}'),
        ('Prepared By:', '{{PREPARED_BY}}'),
        ('Investment Stage:', '{{INVESTMENT_STAGE}}'),
        ('State of Incorporation:', '{{COMPANY_STATE}}'),
        ('Entity Type:', '{{ENTITY_TYPE}}')
    ]
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Executive Summary
    add_heading(doc, 'Executive Summary', 1)
    add_paragraph(doc, '{{EXECUTIVE_SUMMARY}}')

    # Issue Summary
    add_heading(doc, 'Issue Summary', 1)

    add_heading(doc, 'Critical Issues (🔴)', 2)
    add_paragraph(doc, '{{CRITICAL_ISSUES}}')

    add_heading(doc, 'Material Issues (🟠)', 2)
    add_paragraph(doc, '{{MATERIAL_ISSUES}}')

    add_heading(doc, 'Minor Issues (🟡)', 2)
    add_paragraph(doc, '{{MINOR_ISSUES}}')

    # Detailed Findings
    add_heading(doc, 'Detailed Findings', 1)

    sections = [
        ('Corporate Formation & Governance', '{{CORPORATE_FINDINGS}}'),
        ('Capitalization & Securities', '{{CAP_TABLE_FINDINGS}}'),
        ('Intellectual Property', '{{IP_FINDINGS}}'),
        ('Employment & Founders', '{{EMPLOYMENT_FINDINGS}}'),
        ('Material Contracts', '{{CONTRACT_FINDINGS}}'),
        ('Compliance & Regulatory', '{{COMPLIANCE_FINDINGS}}'),
        ('Litigation & Disputes', '{{LITIGATION_FINDINGS}}'),
        ('Financial', '{{FINANCIAL_FINDINGS}}')
    ]

    for section_title, placeholder in sections:
        add_heading(doc, section_title, 2)
        add_paragraph(doc, placeholder)

    # Recommendations
    add_heading(doc, 'Recommendations', 1)
    add_paragraph(doc, '{{RECOMMENDATIONS}}')

    # Open Items
    add_heading(doc, 'Open Items', 1)
    add_paragraph(doc, '{{OPEN_ITEMS}}')

    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer.add_run('DISCLAIMER: ').bold = True
    disclaimer.add_run('This report is provided for informational purposes only and does not constitute legal advice. All findings should be reviewed by qualified legal counsel before any investment decision.')

    doc.save(output_path)
    print(f"Created: {output_path}")


def create_issue_summary(output_path: Path):
    """Create the issue summary template."""
    doc = Document()

    title = doc.add_heading('Due Diligence Issue Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_paragraph(doc, 'Company: {{COMPANY_NAME}}')
    add_paragraph(doc, 'Review Date: {{REVIEW_DATE}}')
    add_paragraph(doc, 'Reviewed By: {{REVIEWER_NAME}}')

    # Summary stats
    doc.add_paragraph()
    add_heading(doc, 'Issue Count', 2)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    data = [
        ('🔴 Critical Issues:', '{{TOTAL_CRITICAL}}'),
        ('🟠 Material Issues:', '{{TOTAL_MATERIAL}}'),
        ('🟡 Minor Issues:', '{{TOTAL_MINOR}}')
    ]
    for i, (label, value) in enumerate(data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Issue list
    doc.add_paragraph()
    add_heading(doc, 'Issues Identified', 2)
    add_paragraph(doc, '{{ISSUE_LIST}}')

    # Next steps
    add_heading(doc, 'Next Steps', 2)
    add_paragraph(doc, '{{NEXT_STEPS}}')

    doc.save(output_path)
    print(f"Created: {output_path}")


def create_cap_table_memo(output_path: Path):
    """Create the cap table memo template."""
    doc = Document()

    title = doc.add_heading('Cap Table Analysis Memorandum', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    add_paragraph(doc, 'Company: {{COMPANY_NAME}}')
    add_paragraph(doc, 'Analysis Date: {{ANALYSIS_DATE}}')
    add_paragraph(doc, 'Prepared By: {{PREPARED_BY}}')

    # Summary
    add_heading(doc, 'Capitalization Summary', 1)

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ('Founder Equity:', '{{FOUNDERS_EQUITY}}'),
        ('Option Pool Size:', '{{OPTION_POOL_SIZE}}'),
        ('Option Pool Available:', '{{OPTION_POOL_AVAILABLE}}'),
        ('SAFE Notes Total:', '{{SAFE_NOTES_TOTAL}}'),
        ('Convertible Notes Total:', '{{CONVERTIBLE_NOTES_TOTAL}}'),
        ('Fully Diluted Shares:', '{{FULLY_DILUTED_SHARES}}')
    ]
    for i, (label, value) in enumerate(data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Pro forma
    add_heading(doc, 'Pro Forma Ownership (Post-Financing)', 1)
    add_paragraph(doc, '{{PRO_FORMA_OWNERSHIP}}')

    # Issues
    add_heading(doc, 'Cap Table Issues Identified', 1)
    add_paragraph(doc, '{{CAP_TABLE_ISSUES}}')

    # Recommendations
    add_heading(doc, 'Recommendations', 1)
    add_paragraph(doc, '{{RECOMMENDATIONS}}')

    doc.save(output_path)
    print(f"Created: {output_path}")


def main():
    """Create all placeholder templates."""
    templates_dir = Path(__file__).parent.parent / 'assets' / 'templates'
    templates_dir.mkdir(parents=True, exist_ok=True)

    print("Creating placeholder templates...")
    print(f"Output directory: {templates_dir}")
    print()

    create_document_request_list(templates_dir / 'document-request-list.docx')
    create_diligence_report(templates_dir / 'diligence-report.docx')
    create_issue_summary(templates_dir / 'issue-summary.docx')
    create_cap_table_memo(templates_dir / 'cap-table-memo.docx')

    print()
    print("Done! Templates created with {{PLACEHOLDER}} syntax.")
    print("Replace these with professionally designed templates as needed.")


if __name__ == "__main__":
    main()
